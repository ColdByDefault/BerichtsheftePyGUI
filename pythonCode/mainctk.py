

# importrs
import tkinter as tk
import customtkinter as ctk
from tkinter import messagebox
from docx import Document
from datetime import datetime
from notes import notes
import requests
import json
from pathlib import Path


def update_notes():
    local_notes_path = Path('pythonCode\\notes.py')  # Adjust the path as needed
    data_url = 'https://raw.githubusercontent.com/ColdByDefault/BerichtsheftePyGUI/master/pythonCode/notes.py'  # URL to check for updated data
    
    try:
        response = requests.get(data_url)
        response.raise_for_status()  # Raises an HTTPError if the response was unsuccessful
        new_notes = response.json()

        # Assuming you're updating a notes.py Python file directly, 
        # convert the JSON back into a Python dictionary format and overwrite notes.py
        with open(local_notes_path, 'w') as file:
            file.write(f"notes = {json.dumps(new_notes, indent=4)}")

        print("Notes updated successfully.")
    except Exception as e:
        print(f"Failed to update notes: {e}")

# Call the update function at an appropriate place in your application
update_notes()

# Create the main window using CustomTkinter's CTk class
root = ctk.CTk()
root.title("Berichtsheft GFN/HD 2024")

# Apply a CustomTkinter theme
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")

root.geometry("600x800")  # Adjust window size as needed
root.resizable(False, False)  # Window non-resizable

# Function to switch frames
def raise_frame(frame):
    frame.tkraise()

# Create frames for each tab content using CustomTkinter CTkFrame
frame1 = ctk.CTkFrame(root)
frame2 = ctk.CTkFrame(root)

# Position the frames in the same grid location; they will be stacked
for frame in (frame1, frame2):
    frame.grid(row=1, column=0, columnspan=4, sticky='nsew', padx=20, pady=20)

root.columnconfigure(0, weight=1)
root.columnconfigure(1, weight=1)
root.columnconfigure(2, weight=1)
root.columnconfigure(3, weight=1)
root.rowconfigure(1, weight=1)

# Tab buttons to switch frames
tab1_button = ctk.CTkButton(root, text='Allgemeine Information', command=lambda: raise_frame(frame1), width=200, corner_radius=10, hover=True)
tab2_button = ctk.CTkButton(root, text='Tägliche Notes', command=lambda: raise_frame(frame2), width=200, corner_radius=10, hover=True)

tab1_button.grid(row=0, column=0, pady=8, padx=50)
tab2_button.grid(row=0, column=2, pady=8, padx=50)

def update_text_fields():
    selected_lf = replacements['[lf_num]'].get()
    schedule = notes.get(selected_lf, {})
    for i, day in enumerate(['montag', 'dienstag', 'mittwoch', 'donnerstag', 'freitag']):
        text_boxes[i].delete("1.0", "end")
        text_boxes[i].insert("1.0", schedule.get(day, "Keine Information"))

# Function to find and replace text in the Word document
def find_and_replace(document, replacements):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        placeholder = f'[{key.strip("[]")}]'
                        if placeholder in paragraph.text:
                            print(f"Replacing '{placeholder}' with '{value.get()}'")
                            paragraph.text = paragraph.text.replace(placeholder, value.get())

# Function to collect text from tk.Text and update replacements
def update_replacements_with_text():
    daily_vars = ['[text1]', '[text2]', '[text3]', '[text4]', '[text5]']
    for var, text_box in zip(daily_vars, text_boxes):
        text_content = text_box.get("1.0", "end-1c")  # Get text content
        replacements[var].set(text_content)  # Update the corresponding StringVar
        
# Function to save the updated document
def save_updated_document():
    try:
        update_replacements_with_text()  # Update replacements dictionary 
        
        selected_lf_num = replacements['[lf_num]'].get()
        current_date = datetime.now().strftime("%d%m%Y")
        file_name = f"{current_date}{selected_lf_num}.docx"
        doc = Document('pythonCode\\Berichtshefte.docx')
        
        find_and_replace(doc, replacements)
        
        doc.save(file_name)
        messagebox.showinfo("Success", "Document saved successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Dictionary for replacements in the Word document
replacements = {
    '[name]': tk.StringVar(value=''),
    '[standort]': tk.StringVar(value=''),
    '[id_num]': tk.StringVar(value=''),
    '[jahr]': tk.StringVar(value=''),
    '[trainer]': tk.StringVar(value=''),
    '[zeitraum_start]': tk.StringVar(value=''),
    '[zeitraum_end]': tk.StringVar(value=''),
    '[lf_num]': tk.StringVar(value=''),  # This will be set by the option menu
    '[text1]': tk.StringVar(value=''),
    '[text2]': tk.StringVar(value=''),
    '[text3]': tk.StringVar(value=''),
    '[text4]': tk.StringVar(value=''),
    '[text5]': tk.StringVar(value=''),
}

# Replace labels and entries in frame1 
info_labels = ['[name]', '[standort]', '[id_num]', '[jahr]', '[trainer]', '[zeitraum_start]', '[zeitraum_end]'] # Placeholders in word
new_labels = ['Vor -Nachname:', 'Standort:', 'Ausbildungsnachweis Nr.:', 'Ausbildungsjahr:', 'Trainer/Dozent:', 'Datum von:', 'Bis:'] # GUI

# looü and create entries
for i, label in enumerate(info_labels, start=0):
    ctk.CTkLabel(frame1, text=new_labels[i]).grid(row=i, column=0, padx=10, pady=3)
    ctk.CTkEntry(frame1, textvariable=replacements[label], width=200, height=55).grid(row=i, column=1, padx=60, pady=8)

# Option menu for [lf_num] with updated functionality to include text field updates
lf_num_options = list(notes.keys())  # Dynamically get LF options from notes
replacements['[lf_num]'].set(lf_num_options[0])  # Set default value
lf_num_menu = ctk.CTkOptionMenu(frame1, variable=replacements['[lf_num]'], values=lf_num_options, command=lambda _: update_text_fields())
lf_num_menu.grid(row=8, column=1, padx=60, pady=8)

# Update Texts Button (new)
update_texts_btn = ctk.CTkButton(frame1, text="Update Texts", command=update_text_fields)
update_texts_btn.grid(row=9, column=1, padx=60, pady=8)

# Text boxes for daily notes in frame2
text_boxes = []
daily_labels = ['Montag:', 'Dienstag:', 'Mittwoch:', 'Donnerstag:', 'Freitag:']
for i, day in enumerate(daily_labels):
    ctk.CTkLabel(frame2, text=day).grid(row=i*2, column=0, sticky='nw', padx=10, pady=(10, 2))
    text_box = ctk.CTkTextbox(frame2, height=60, width=400)
    text_box.grid(row=i*2+1, column=1, padx=10, pady=(10, 10))
    text_boxes.append(text_box)

# btns ctk to save the progress
submit_btn = ctk.CTkButton(root, text="Submit", command=save_updated_document)
submit_btn.grid(row=2, column=0, columnspan=4, pady=20)

# Initialize the first tab
raise_frame(frame1)

root.mainloop()
